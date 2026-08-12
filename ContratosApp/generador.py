"""Generación reutilizable de contratos DOCX, sin dependencia de Streamlit."""

from __future__ import annotations

from dataclasses import dataclass
from datetime import date, datetime
from io import BytesIO
from pathlib import Path
from typing import BinaryIO

from docx import Document

from modelos import DatosContrato
from utils import sanitizar_nombre_archivo


MESES_ES = {
    "January": "enero", "February": "febrero", "March": "marzo",
    "April": "abril", "May": "mayo", "June": "junio",
    "July": "julio", "August": "agosto", "September": "septiembre",
    "October": "octubre", "November": "noviembre", "December": "diciembre",
}

# Se conserva el texto contractual original, aislado para facilitar cambios futuros.
SUFIJO_CONTRACTUAL = "Contrato Apertura de Cuentas_acuerdo 2026"


@dataclass(frozen=True)
class ResultadoContrato:
    contenido: bytes
    nombre_archivo: str
    placeholders_reemplazados: tuple[str, ...]
    placeholders_no_encontrados: tuple[str, ...]


def formatear_fecha(fecha_obj: date | datetime) -> str:
    """Convierte una fecha al formato contractual: 29 del mes de septiembre de 2025."""
    return (
        f"{fecha_obj.strftime('%d')} del mes de "
        f"{MESES_ES[fecha_obj.strftime('%B')]} de {fecha_obj.strftime('%Y')}"
    )


def generar_nombre_archivo(datos: DatosContrato) -> str:
    """Genera el nombre controlado consumido posteriormente por la automatización."""
    fecha_formateada = datos.fecha.strftime("%d%m%Y")
    nombre = (
        f"{fecha_formateada}_{datos.nombres} {datos.apellidos}_"
        f"{datos.banco}_{SUFIJO_CONTRACTUAL}.docx"
    )
    return sanitizar_nombre_archivo(nombre)


def _reemplazar_en_parrafo(parrafo, texto_buscar: str, reemplazo: str) -> int:
    """Reemplaza incluso si un placeholder quedó dividido en varios runs.

    Conserva los runs y sus estilos: el texto nuevo usa el formato del primer run
    afectado y cualquier sufijo conserva el formato del último run afectado.
    """
    reemplazos = 0
    while texto_buscar in parrafo.text:
        texto_completo = "".join(run.text for run in parrafo.runs)
        inicio = texto_completo.find(texto_buscar)
        fin = inicio + len(texto_buscar)
        cursor = 0
        indice_inicio = indice_fin = 0
        offset_inicio = offset_fin = 0

        for indice, run in enumerate(parrafo.runs):
            siguiente = cursor + len(run.text)
            if cursor <= inicio < siguiente:
                indice_inicio, offset_inicio = indice, inicio - cursor
            if cursor < fin <= siguiente:
                indice_fin, offset_fin = indice, fin - cursor
                break
            cursor = siguiente

        runs = parrafo.runs
        if indice_inicio == indice_fin:
            run = runs[indice_inicio]
            run.text = run.text[:offset_inicio] + reemplazo + run.text[offset_fin:]
        else:
            runs[indice_inicio].text = runs[indice_inicio].text[:offset_inicio] + reemplazo
            for indice in range(indice_inicio + 1, indice_fin):
                runs[indice].text = ""
            runs[indice_fin].text = runs[indice_fin].text[offset_fin:]
        reemplazos += 1
    return reemplazos


def reemplazar_texto(elemento, texto_buscar: str, reemplazo: str) -> int:
    """Reemplaza un placeholder en párrafos y tablas, preservando formato."""
    cantidad = 0
    if hasattr(elemento, "paragraphs"):
        for parrafo in elemento.paragraphs:
            cantidad += _reemplazar_en_parrafo(parrafo, texto_buscar, str(reemplazo))
    if hasattr(elemento, "tables"):
        for tabla in elemento.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    cantidad += reemplazar_texto(celda, texto_buscar, reemplazo)
    return cantidad


def _asignar_texto_parrafo(parrafo, texto: str) -> None:
    """Actualiza el contenido de un párrafo sin modificar su estilo de párrafo."""
    if parrafo.runs:
        parrafo.runs[0].text = texto
        for run in parrafo.runs[1:]:
            run.text = ""
    else:
        parrafo.add_run(texto)


def _actualizar_linea_productos(documento, datos: DatosContrato) -> bool:
    """Actualiza la línea fija posterior al encabezado de productos de la plantilla.

    Es un respaldo para las plantillas antiguas que aún contienen, por ejemplo,
    ``Banco BCI: Cuenta Corriente + Tarjeta de crédito`` en lugar de los
    placeholders ``<<Banco>>: <<Productos>>``.
    """
    parrafos = documento.paragraphs
    for indice, parrafo in enumerate(parrafos[:-1]):
        encabezado = parrafo.text.strip().casefold()
        if "los productos entregados en monitoreo son" not in encabezado:
            continue

        for siguiente in parrafos[indice + 1:]:
            texto = siguiente.text.strip()
            if not texto:
                continue
            if texto.casefold().startswith("banco") and ":" in texto:
                _asignar_texto_parrafo(siguiente, f"{datos.banco}: {datos.productos}")
                return True
            break
    return False


def _reemplazos(datos: DatosContrato) -> dict[str, str]:
    return {
        "<<Nombres>>": datos.nombres,
        "<<Apellidos>>": datos.apellidos,
        "<<DNI>>": datos.dni,
        "<<Fecha>>": formatear_fecha(datos.fecha),
        "<<celular>>": datos.celular_contacto,
        "<<email>>": datos.email_personal,
        "<<Ciudad>>": datos.ciudad,
        "<<País>>": datos.pais,
        "<<Monto>>": str(datos.monto),
        "<<Banco>>": datos.banco,
        "<<Productos>>": datos.productos,
    }


def generar_contrato(
    plantilla: str | Path | bytes | BinaryIO, datos: DatosContrato
) -> ResultadoContrato:
    """Genera un DOCX editable en memoria a partir de una plantilla y sus datos."""
    origen = BytesIO(plantilla) if isinstance(plantilla, bytes) else plantilla
    documento = Document(origen)
    encontrados: list[str] = []
    no_encontrados: list[str] = []

    for placeholder, valor in _reemplazos(datos).items():
        cantidad = reemplazar_texto(documento, placeholder, valor)
        # El notebook sólo cubría el cuerpo; estos elementos amplían cobertura sin acoplar UI.
        for seccion in documento.sections:
            cantidad += reemplazar_texto(seccion.header, placeholder, valor)
            cantidad += reemplazar_texto(seccion.footer, placeholder, valor)
        (encontrados if cantidad else no_encontrados).append(placeholder)

    _actualizar_linea_productos(documento, datos)

    salida = BytesIO()
    documento.save(salida)
    return ResultadoContrato(
        contenido=salida.getvalue(),
        nombre_archivo=generar_nombre_archivo(datos),
        placeholders_reemplazados=tuple(encontrados),
        placeholders_no_encontrados=tuple(no_encontrados),
    )
