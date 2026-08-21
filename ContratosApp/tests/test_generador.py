from datetime import date
from io import BytesIO
import unittest

from docx import Document

from generador import formatear_fecha, generar_contrato, generar_nombre_archivo
from modelos import DatosContrato
from utils import parsear_datos_pegados, sanitizar_nombre_archivo


def datos_prueba() -> DatosContrato:
    return DatosContrato(
        nombres="Ana", apellidos="Pérez", dni="123", celular_contacto="", email_personal="ana@example.com",
        ciudad="Santiago", pais="Chile", monto="350", banco="Banco/Estado", productos="Cuenta", fecha=date(2027, 1, 2),
    )


class GeneradorTests(unittest.TestCase):
    def test_fecha_dinamica(self):
        self.assertEqual(formatear_fecha(date(2027, 1, 2)), "02 del mes de enero de 2027")
        self.assertTrue(generar_nombre_archivo(datos_prueba()).startswith("02012027_"))

    def test_sanitizar_nombre(self):
        self.assertEqual(sanitizar_nombre_archivo('A<B>:C"/D\\E|F?G*'), "ABCDEFG")

    def test_nombre_archivo(self):
        nombre = generar_nombre_archivo(datos_prueba())
        self.assertEqual(nombre, "02012027_Ana Pérez_BancoEstado_Contrato Apertura de Cuentas_acuerdo 2026.docx")

    def test_reemplaza_placeholder_dividido_en_runs_y_tabla(self):
        documento = Document()
        parrafo = documento.add_paragraph()
        parrafo.add_run("Cliente: <<Nom")
        parrafo.add_run("bres>>")
        celda = documento.add_table(rows=1, cols=1).cell(0, 0)
        celda.text = "Banco: <<Banco>>"
        plantilla = BytesIO()
        documento.save(plantilla)

        resultado = generar_contrato(plantilla.getvalue(), datos_prueba())
        generado = Document(BytesIO(resultado.contenido))
        self.assertEqual(generado.paragraphs[0].text, "Cliente: Ana")
        self.assertEqual(generado.tables[0].cell(0, 0).text, "Banco: Banco/Estado")

    def test_actualiza_linea_fija_de_banco_y_productos(self):
        documento = Document()
        documento.add_paragraph("Los productos entregados en monitoreo son:")
        documento.add_paragraph("Banco BCI: Cuenta Corriente + Tarjeta de crédito")
        plantilla = BytesIO()
        documento.save(plantilla)

        resultado = generar_contrato(plantilla.getvalue(), datos_prueba())
        generado = Document(BytesIO(resultado.contenido))
        self.assertEqual(generado.paragraphs[1].text, "Banco/Estado: Cuenta")

    def test_clasifica_datos_pegados_desde_excel(self):
        texto = "Nombres: Ana María\nEmail personal: ana@example.com\nPaís: España\nCelular de Contacto: 34 664"
        campos, no_reconocidas = parsear_datos_pegados(texto)
        self.assertEqual(campos["nombres"], "Ana María")
        self.assertEqual(campos["email"], "ana@example.com")
        self.assertEqual(campos["pais"], "España")
        self.assertEqual(campos["celular"], "34 664")
        self.assertEqual(no_reconocidas, [])

    def test_clasifica_fila_de_excel_con_encabezados(self):
        texto = "Nombres\tApellidos\tDNI\tEmail\tBanco\nAna\tDemo\t123\tana@ejemplo.test\tBanco Prueba"
        campos, no_reconocidas = parsear_datos_pegados(texto)
        self.assertEqual(no_reconocidas, [])
        self.assertEqual(campos["nombres"], "Ana")
        self.assertEqual(campos["apellidos"], "Demo")
        self.assertEqual(campos["email"], "ana@ejemplo.test")

    def test_limpia_texto_libre_sin_etiquetas(self):
        texto = "ANA MARIA\nPEREZ SOTO\n12.345.678-9\n+56 9 1234 5678\nContacto <ANA@EJEMPLO.COM>\nSANTIAGO\nCHILE"
        campos, no_reconocidas = parsear_datos_pegados(texto)
        self.assertEqual(no_reconocidas, [])
        self.assertEqual(campos["nombres"], "Ana Maria")
        self.assertEqual(campos["dni"], "12.345.678-9")
        self.assertEqual(campos["celular"], "+56912345678")
        self.assertEqual(campos["email"], "ana@ejemplo.com")

    def test_clasifica_json_de_contacto(self):
        texto = '{"nombre_completo":"Ana Maria Perez Soto","rut":"12.345.678-9","telefono":"+56 9 1234 5678","correo":"ANA@EJEMPLO.COM","ciudad":"santiago","pais":"chile"}'
        campos, no_reconocidas = parsear_datos_pegados(texto)
        self.assertEqual(no_reconocidas, [])
        self.assertEqual(campos["nombres"], "Ana Maria")
        self.assertEqual(campos["apellidos"], "Perez Soto")
        self.assertEqual(campos["email"], "ana@ejemplo.com")


if __name__ == "__main__":
    unittest.main()
